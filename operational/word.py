import pandas as pd
from docx import Document

csv_file = '../results/results.csv'
df = pd.read_csv(csv_file)
doc = Document()
doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
header_row = doc.tables[0].rows[0].cells

for col_idx, column in enumerate(df.columns):
    header_row[col_idx].text = str(column)

for _, row in df.iterrows():
    new_row = doc.tables[0].add_row().cells
    for col_idx, value in enumerate(row):
        new_row[col_idx].text = str(value)

doc.save('output_document.docx')
