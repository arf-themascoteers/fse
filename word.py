import pandas as pd
from docx import Document

# Load the CSV file into a Pandas DataFrame
csv_file = 'results/1700307469.csv'  # Replace with your actual CSV file
df = pd.read_csv(csv_file)

# Create a new Word document
doc = Document()

# Add a table to the Word document
doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')

# Access the first row of the table to add column names
header_row = doc.tables[0].rows[0].cells
for col_idx, column in enumerate(df.columns):
    header_row[col_idx].text = str(column)

# Add data rows to the table
for _, row in df.iterrows():
    new_row = doc.tables[0].add_row().cells
    for col_idx, value in enumerate(row):
        new_row[col_idx].text = str(value)

# Save the Word document
doc.save('output_document.docx')  # Replace with your desired output file name
